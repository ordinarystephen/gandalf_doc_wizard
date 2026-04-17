"""Generate test fixture files for ingest tests."""
import os
os.makedirs("tests/fixtures", exist_ok=True)

def make_sample_pdf(path="tests/fixtures/sample.pdf"):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(path, pagesize=letter)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 720, "FINANCIAL ANALYSIS")
    c.setFont("Helvetica", 11)
    c.drawString(72, 700, "The borrower demonstrates strong cash flow coverage.")
    c.drawString(72, 685, "DSCR is 1.35x based on trailing twelve months.")
    c.drawString(72, 670, "The covenant threshold is 1.20x per the credit agreement.")
    c.save()

def make_sample_docx(path="tests/fixtures/sample.docx"):
    from docx import Document
    doc = Document()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("This loan supports the acquisition of an industrial property.")
    doc.add_heading("Financial Analysis", level=1)
    doc.add_paragraph("The DSCR covenant threshold is 1.20x per the credit agreement.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "LTV"
    table.cell(1, 1).text = "65%"
    doc.save(path)

def make_sample_xlsx(path="tests/fixtures/sample.xlsx"):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Financials"
    ws.append(["Metric", "2023", "2024"])
    ws.append(["Revenue", 1000000, 1200000])
    ws.append(["EBITDA", 300000, 380000])
    ws2 = wb.create_sheet("Covenants")
    ws2.append(["Covenant", "Threshold", "Actual"])
    ws2.append(["DSCR", "1.20x", "1.35x"])
    wb.save(path)

if __name__ == "__main__":
    make_sample_pdf()
    make_sample_docx()
    make_sample_xlsx()
    print("Fixtures created.")
